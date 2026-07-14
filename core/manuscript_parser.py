"""워드(.docx) 원고 파서.

워드 중간의 구분 표(순번/이름/URL/제목 2열 표)를 기준으로 여러 명의 원고를 분리한다.
취소선(빨간 줄) 텍스트는 '삭제 표시'로 보고 본문에서 제외하되 개수를 세어 알려준다.
parse_docx_sections(file_bytes) -> [{order, name, url, title, body, deleted}]
all_text(file_bytes) -> str   (구분 표가 없을 때 전체 본문, 취소선 제외)
"""
from __future__ import annotations
import io
import re
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

_LABELS = ("순번", "이름", "URL", "제목")


def _run_struck(r) -> bool:
    """런에 취소선(strike/dstrike)이 켜져 있는지."""
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        return False
    for tag in ("w:strike", "w:dstrike"):
        el = rpr.find(qn(tag))
        if el is not None and el.get(qn("w:val")) not in ("false", "0", "none"):
            return True
    return False


def _para_parts(paragraph: Paragraph) -> tuple[str, str]:
    """문단을 (보이는 텍스트, 취소선=삭제 텍스트)로 분리.
    하이퍼링크 안의 텍스트와 실제 URL(관계 대상)까지 포함한다
    (python-docx 의 paragraph.runs 는 하이퍼링크 런을 누락하므로 XML 직접 순회)."""
    visible, struck = [], []
    for r in paragraph._p.iter(qn("w:r")):
        text = "".join(t.text or "" for t in r.findall(qn("w:t")))
        if not text:
            continue
        (struck if _run_struck(r) else visible).append(text)
    vis = "".join(visible)
    # 하이퍼링크 실제 URL 보강 (표시 텍스트가 URL이 아니어도 캡처)
    for h in paragraph._p.iter(qn("w:hyperlink")):
        rid = h.get(qn("r:id"))
        if not rid:
            continue
        try:
            url = paragraph.part.rels[rid].target_ref
        except (KeyError, AttributeError):
            url = None
        if url and url not in vis:
            vis = (vis + " " + url).strip()
    return vis, "".join(struck)


def _divider_info(tbl: Table) -> dict | None:
    """표가 '순번/이름/URL/제목' 구분 표면 그 값을 반환, 아니면 None."""
    found = {}
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = cells[0].text.strip()
            if key in _LABELS:
                found[key] = cells[1].text.strip()
    if "이름" in found or "제목" in found:
        return {"order": found.get("순번", ""), "name": found.get("이름", ""),
                "url": found.get("URL", ""), "title": found.get("제목", "")}
    return None


def _blips_blobs(element, rels) -> list[bytes]:
    """XML 요소(문단/표) 안의 이미지(a:blip)들을 등장 순서대로 bytes 로."""
    out = []
    for blip in element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid and rid in rels:
            blob = getattr(rels[rid], "blob", None)
            if blob:
                out.append(blob)
    return out


def parse_docx_sections(file_bytes: bytes) -> list[dict]:
    doc = Document(io.BytesIO(file_bytes))
    rels = doc.part.related_parts
    sections: list[dict] = []
    current: dict | None = None
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:tbl"):
            tbl = Table(child, doc)
            info = _divider_info(tbl)
            if info is not None:
                current = {**info, "_body": [], "_deleted": [], "_images": []}
                sections.append(current)
                continue
            if current is not None:  # 일반 표 → 본문에 셀 텍스트 포함
                for row in tbl.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            current["_body"].append(t)
                current["_images"].extend(_blips_blobs(child, rels))
        elif child.tag == qn("w:p"):
            if current is not None:
                vis, struck = _para_parts(Paragraph(child, doc))
                if vis.strip():
                    current["_body"].append(vis.strip())
                if struck.strip():
                    current["_deleted"].append(struck.strip())
                current["_images"].extend(_blips_blobs(child, rels))
    out = []
    for s in sections:
        out.append({"order": s["order"], "name": s["name"], "url": s["url"],
                    "title": s["title"], "body": "\n".join(s["_body"]),
                    "deleted": s["_deleted"], "images": s["_images"]})
    return out


def all_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    out = []
    for p in doc.paragraphs:
        vis, _ = _para_parts(p)
        if vis.strip():
            out.append(vis.strip())
    return "\n".join(out)


def paragraph_pages(file_bytes: bytes) -> list[tuple[int, str]]:
    """문단별 (추정 페이지, 텍스트). Word 의 lastRenderedPageBreak/페이지나눔 기준(추정).
    Word 로 저장된 문서만 페이지 경계가 기록되며, 없으면 전부 1쪽으로 나온다."""
    doc = Document(io.BytesIO(file_bytes))
    out, page = [], 1
    for para in doc.paragraphs:
        vis, _ = _para_parts(para)
        out.append((page, vis))
        p = para._p
        brk = len(p.findall(".//" + qn("w:lastRenderedPageBreak")))
        brk += sum(1 for b in p.findall(".//" + qn("w:br")) if b.get(qn("w:type")) == "page")
        page += brk
    return out


def find_page(pages: list[tuple[int, str]], needle: str) -> int | None:
    """needle(공백 무시)이 처음 등장하는 문단의 추정 페이지."""
    if not needle or not pages:
        return None
    key = re.sub(r"\s+", "", needle)
    if not key:
        return None
    for page, text in pages:
        if key in re.sub(r"\s+", "", text):
            return page
    return None


def extract_images(file_bytes: bytes) -> list[bytes]:
    """워드에 삽입된 이미지들을 문서 등장 순서대로 bytes 리스트로 반환.
    (여러 명 원고가 든 문서면 문서 전체 이미지가 나온다 — 블로거별 비교는 해당 원고만 올려서.)"""
    doc = Document(io.BytesIO(file_bytes))
    rels = doc.part.related_parts
    out = []
    for blip in doc.element.body.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid or rid not in rels:
            continue
        blob = getattr(rels[rid], "blob", None)
        if blob:
            out.append(blob)
    return out


def read_pdf(file_bytes: bytes) -> str:
    """PDF 전체 텍스트 추출(pypdf). 스캔본(이미지 PDF)은 텍스트가 거의 안 나올 수 있음."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages)
