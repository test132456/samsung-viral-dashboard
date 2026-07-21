import io
from docx import Document
from core import manuscript_parser


def _make_docx(entries):
    """entries: [(순번, 이름, url, 제목, [본문줄...])] → docx bytes."""
    d = Document()
    for order, name, url, title, body_lines in entries:
        t = d.add_table(rows=4, cols=2)
        t.cell(0, 0).text, t.cell(0, 1).text = "순번", order
        t.cell(1, 0).text, t.cell(1, 1).text = "이름", name
        t.cell(2, 0).text, t.cell(2, 1).text = "URL", url
        t.cell(3, 0).text, t.cell(3, 1).text = "제목", title
        for line in body_lines:
            d.add_paragraph(line)
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def test_splits_by_blogger():
    data = _make_docx([
        ("(1)", "제아", "https://blog.naver.com/jea", "제아 제목", ["제아 본문 첫 줄", "둘째 줄"]),
        ("(2)", "여름", "https://blog.naver.com/hylim1224", "여름 제목입니다", ["여름 본문 첫 줄"]),
    ])
    secs = manuscript_parser.parse_docx_sections(data)
    assert len(secs) == 2
    assert [s["name"] for s in secs] == ["제아", "여름"]
    summer = [s for s in secs if s["name"] == "여름"][0]
    assert summer["title"] == "여름 제목입니다"
    assert summer["url"] == "https://blog.naver.com/hylim1224"
    assert "여름 본문 첫 줄" in summer["body"]
    assert "제아 본문" not in summer["body"]   # 다른 사람 본문 섞이지 않음


def test_no_divider_returns_empty_sections():
    d = Document()
    d.add_paragraph("구분 표 없는 일반 원고")
    b = io.BytesIO(); d.save(b)
    assert manuscript_parser.parse_docx_sections(b.getvalue()) == []
    assert "일반 원고" in manuscript_parser.all_text(b.getvalue())


def test_strikethrough_excluded_and_counted():
    d = Document()
    t = d.add_table(rows=4, cols=2)
    t.cell(0, 0).text, t.cell(0, 1).text = "순번", "(1)"
    t.cell(1, 0).text, t.cell(1, 1).text = "이름", "여름"
    t.cell(2, 0).text, t.cell(2, 1).text = "URL", "https://x"
    t.cell(3, 0).text, t.cell(3, 1).text = "제목", "제목입니다"
    p = d.add_paragraph()
    p.add_run("정상 문장입니다 ")
    dele = p.add_run("삭제된 부분입니다")
    dele.font.strike = True
    b = io.BytesIO(); d.save(b)
    secs = manuscript_parser.parse_docx_sections(b.getvalue())
    assert len(secs) == 1
    assert "정상 문장입니다" in secs[0]["body"]
    assert "삭제된 부분" not in secs[0]["body"]        # 취소선 제외
    assert secs[0]["deleted"] and "삭제된 부분입니다" in secs[0]["deleted"][0]


def _add_hyperlink(paragraph, url, text):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rid = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = text
    r.append(t); hl.append(r); paragraph._p.append(hl)


def test_hyperlink_text_and_url_extracted():
    d = Document()
    p = d.add_paragraph("자세한 내용은 ")
    _add_hyperlink(p, "https://direct.samsungfire.com/mall/PP030701_001.html", "여기")
    b = io.BytesIO(); d.save(b)
    txt = manuscript_parser.all_text(b.getvalue())
    assert "여기" in txt                                                   # 표시 텍스트
    assert "direct.samsungfire.com/mall/PP030701_001.html" in txt          # 실제 URL


def test_find_page_returns_page():
    d = Document()
    d.add_paragraph("첫 페이지 내용")
    b = io.BytesIO(); d.save(b)
    pages = manuscript_parser.paragraph_pages(b.getvalue())
    assert manuscript_parser.find_page(pages, "첫 페이지") == 1
    assert manuscript_parser.find_page(pages, "없는문구") is None


def test_find_page_ignores_punctuation_hashtag_and_multi_sentence():
    # 원문 문단엔 마침표·해시태그가 있고, 비교문장은 여러 문장이 마침표 없이 합쳐진 형태
    pages = [(1, "국젯선 유류할증료가 인하되고 일봉 여행 가는 사람들이 늘어나고 있어요. #해외여행보험"),
             (2, "여행자 보험을 팻스하는 경우가 많은데요.")]
    n_multi = "국젯선 유류할증료가 인하되고 일봉 여행 가는 사람들이 늘어나고 있어요 해외여행보험"
    assert manuscript_parser.find_page(pages, n_multi) == 1     # 마침표/해시태그 무시하고 매칭
    assert manuscript_parser.find_page(pages, "여행자 보험을 팻스하는 경우가 많은데요") == 2


def test_find_page_context_spanning_paragraphs():
    # 심의 표현 점검의 문맥 조각처럼 여러 문단(줄)에 걸친 needle 도 위치로 페이지를 찾는다
    pages = [(9, "비행기 결항"), (9, "캐리어 파손 등"), (9, "다양한 일들을 겪고")]
    assert manuscript_parser.find_page(pages, "…결항, 캐리어 파손 등 다양한 일들…") == 9
    pages2 = [(4, "앞 문단 내용"), (5, "막 먹었더니 배가"), (5, "아파서 병원에 갔죠")]
    assert manuscript_parser.find_page(pages2, "먹었더니 배가 아파서 병원") == 5
