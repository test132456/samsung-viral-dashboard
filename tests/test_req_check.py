import io
import struct
import zlib
from docx import Document
from core import req_check, manuscript_parser


def _png(w=2, h=2) -> bytes:
    """유효한 최소 PNG(RGB) 생성 — docx add_picture 용."""
    def chunk(typ, data):
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xffffffff)
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
    raw = (b"\x00" + b"\xff\x00\x00" * w) * h
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


_PNG = _png()


def _full_body():
    return ("*이 포스팅은 삼성화재 다이렉트로부터 소정의 광고비(원고료)를 받아 작성되었습니다\n"
            "삼성화재 다이렉트 해외여행보험 소개입니다. 여행자보험과 비교해 장점이 큽니다. "
            "해외여행보험 가입 팁과 해외여행보험 추천 이유를 정리했어요.\n"
            "가입 링크 https://direct.samsungfire.com/mall/PP030701_001.html\n"
            "준법감시인확인필 제00호\n"
            "#삼성화재다이렉트 #삼성화재다이렉트해외여행보험 #해외여행보험 #항공지연 "
            "#지연결항 #해외여행자보험 #해외여행보험추천 #해외여행보험후기")


def test_full_manuscript_passes():
    title = "해외여행보험 출국 귀국 항공기 지연 보상 추천"
    items = req_check.evaluate(title, _full_body(), image_count=12, is_official=False)
    by = {i["name"]: i for i in items}
    assert by["브랜드·보험명 정확 표기"]["status"] == "ok"
    assert by["타 보험사 언급 지양"]["status"] == "ok"
    assert by["필수 해시태그 포함"]["status"] == "ok"
    assert by["이미지 10개 이상"]["status"] == "ok"
    assert by["허용 상품 링크만 사용"]["status"] == "ok"
    assert by["유료광고 문안(상단)"]["status"] == "ok"
    assert by["준법감시인확인필(하단)"]["status"] == "ok"
    assert "<b>" in by["브랜드·보험명 정확 표기"]["evidence"]   # 근거에 매칭어 강조


def test_competitor_mention_flagged():
    items = req_check.evaluate("해외여행보험", "현대해상 여행자보험도 좋아요", image_count=0)
    by = {i["name"]: i for i in items}
    assert by["타 보험사 언급 지양"]["status"] == "fail"
    assert "현대해상" in by["타 보험사 언급 지양"]["detail"]


def test_foreign_link_flagged():
    items = req_check.evaluate("해외여행보험", "링크 https://hi.example.com/x 참고", image_count=0)
    by = {i["name"]: i for i in items}
    assert by["허용 상품 링크만 사용"]["status"] == "fail"


def test_keyword_count_excludes_hashtags():
    body = "해외여행보험 소개 #해외여행보험 #해외여행보험추천 #해외여행보험후기"  # 본문 1개
    items = req_check.evaluate("t", body, image_count=0)
    by = {i["name"]: i for i in items}
    assert "1개" in by["'해외여행보험' 키워드 3~5개"]["detail"]


def test_official_blog_paid_ad_na():
    items = req_check.evaluate("해외여행보험", "본문", image_count=0, is_official=True)
    by = {i["name"]: i for i in items}
    assert by["유료광고 문안(상단)"]["status"] == "na"


def test_count_images():
    d = Document(); d.add_paragraph("no image")
    b = io.BytesIO(); d.save(b)
    assert manuscript_parser.count_images(b.getvalue()) == 0
    d2 = Document(); d2.add_picture(io.BytesIO(_PNG))
    b2 = io.BytesIO(); d2.save(b2)
    assert manuscript_parser.count_images(b2.getvalue()) >= 1
