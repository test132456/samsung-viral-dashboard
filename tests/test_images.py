import io
import struct
import zlib
from docx import Document
from core import manuscript_parser, fetcher


def _png(w=3, h=3) -> bytes:
    def chunk(typ, data):
        b = typ + data
        return struct.pack(">I", len(data)) + b + struct.pack(">I", zlib.crc32(b) & 0xffffffff)
    raw = (b"\x00" + b"\x10\x20\x30" * w) * h
    return (b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b""))


def test_extract_images_from_docx_in_order():
    d = Document()
    d.add_paragraph("첫 문단")
    d.add_picture(io.BytesIO(_png()))
    d.add_paragraph("둘째 문단")
    d.add_picture(io.BytesIO(_png(4, 4)))
    b = io.BytesIO(); d.save(b)
    imgs = manuscript_parser.extract_images(b.getvalue())
    assert len(imgs) == 2
    assert all(isinstance(x, bytes) and x[:8] == b"\x89PNG\r\n\x1a\n" for x in imgs)


def test_extract_images_none():
    d = Document(); d.add_paragraph("이미지 없음")
    b = io.BytesIO(); d.save(b)
    assert manuscript_parser.extract_images(b.getvalue()) == []


def test_extract_image_urls_order_and_lazy():
    html = """
    <div class="se-main-container">
      <img src="https://postfiles.pstatic.net/a.jpg?type=w966">
      <img data-lazy-src="https://postfiles.pstatic.net/b.jpg" src="data:image/gif;base64,zzz">
      <img src="https://ssl.pstatic.net/static/ui-icon.png">
      <img src="https://postfiles.pstatic.net/a.jpg?type=w80">
    </div>"""
    urls = fetcher.extract_image_urls(html)
    assert urls[0].endswith("a.jpg?type=w966")
    assert urls[1] == "https://postfiles.pstatic.net/b.jpg"   # data-lazy-src 우선
    assert not any("ui-icon" in u for u in urls)              # UI 자산 제외
    assert len(urls) == 2                                     # a.jpg 중복 제거


def _divider(doc, name):
    t = doc.add_table(rows=4, cols=2)
    t.cell(0, 0).text, t.cell(0, 1).text = "순번", "(1)"
    t.cell(1, 0).text, t.cell(1, 1).text = "이름", name
    t.cell(2, 0).text, t.cell(2, 1).text = "URL", "https://x"
    t.cell(3, 0).text, t.cell(3, 1).text = "제목", f"{name} 제목"


def test_images_attributed_per_blogger_section():
    # 제아 구간 1장, 여름 구간 2장 → 각 구간에 제대로 귀속되는지
    d = Document()
    _divider(d, "제아")
    d.add_paragraph("제아 본문")
    d.add_picture(io.BytesIO(_png()))                # 제아 이미지 1
    _divider(d, "여름")
    d.add_paragraph("여름 본문")
    d.add_picture(io.BytesIO(_png(4, 4)))            # 여름 이미지 1
    d.add_picture(io.BytesIO(_png(5, 5)))            # 여름 이미지 2
    b = io.BytesIO(); d.save(b)
    secs = manuscript_parser.parse_docx_sections(b.getvalue())
    by = {s["name"]: s for s in secs}
    assert len(by["제아"]["images"]) == 1
    assert len(by["여름"]["images"]) == 2
    assert all(x[:8] == b"\x89PNG\r\n\x1a\n" for x in by["여름"]["images"])


def test_extract_image_urls_excludes_stickers():
    from core import fetcher
    html = """
    <div class="se-main-container">
      <div class="se-component se-image"><img src="https://postfiles.pstatic.net/photo1.jpg?type=w966"></div>
      <div class="se-component se-sticker"><img src="https://storep-phinf.pstatic.net/ogq_x.png"></div>
      <div class="se-component se-image"><img src="https://postfiles.pstatic.net/photo2.jpg"></div>
      <img class="se-sticker-image" src="https://postfiles.pstatic.net/looks-like-photo-but-sticker.png">
    </div>"""
    urls = fetcher.extract_image_urls(html)
    assert urls == ["https://postfiles.pstatic.net/photo1.jpg?type=w966",
                    "https://postfiles.pstatic.net/photo2.jpg"]   # 사진 2장만, 스티커 제외
